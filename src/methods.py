### `src/methods.py`
### Main application for the RAG on Snow project
### Open-Source, hosted on https://github.com/DrBenjamin/Dateiablage
### Please reach out to ben@seriousbenentertainment.org for any questions
## Modules
import wx
import os
import sys
import subprocess
import webbrowser
import src.globals as g
from docx import Document
from src.files import list_files
from src.tasks import on_import_tasks
from src.learning import display_learning

# Method to handle the right click event
def on_right_click(self, event):
    # Creating the context menu
    menu = wx.Menu()
    open_item = menu.Append(wx.ID_ANY, "Öffnen")
    copy_path = menu.Append(wx.ID_ANY, "Kopiere Pfad")
    convert_item = menu.Append(wx.ID_ANY, "Konvertiere srt in vtt")

    # Binding handlers
    self.Bind(wx.EVT_MENU, self.on_file_activated, open_item)
    self.Bind(wx.EVT_MENU, self.on_copy_path, copy_path)
    self.Bind(wx.EVT_MENU, self.on_convert, convert_item)

    # Showing the menu
    self.PopupMenu(menu, event.GetPosition())
    menu.Destroy()

# Method to handle the Copy Path menu item
def on_copy_path(self, event):
    if g.file_path is not None:
        wx.TheClipboard.Open()
        wx.TheClipboard.SetData(wx.TextDataObject(g.file_path))
        wx.TheClipboard.Close()
    else:
        wx.MessageBox("Keine Datei ausgewählt", "Error", wx.OK | wx.ICON_ERROR)

# Method to handle the Convert menu item
def on_convert(self, event):
    try:
        if g.file_path.endswith(".srt"):
            convert_srt_to_vtt(
                    g.file_path,
                    overwrite = self.config.ReadBool("srt_converter_overwrite")
                )
            self.on_refresh(None)
        else:
            wx.MessageBox("Dateiendung wird nicht unterstützt", "Error", wx.OK | wx.ICON_ERROR)
    except Exception as e:
        wx.MessageBox(f"Datei konnte nicht konvertiert werden: {e}", "Error", wx.OK | wx.ICON_ERROR)

# Method to convert subtitles (srt into vtt)
def convert_srt_to_vtt(file_path, overwrite = False):
    vtt_file_path = file_path.rsplit(".", 1)[0] + ".vtt"

    if os.path.exists(vtt_file_path):
        file_name = os.path.basename(vtt_file_path)
        if not overwrite:
            response = wx.MessageBox(
                f"Die Datei '{file_name}' existiert bereits. Möchten Sie sie ersetzen?",
                "Datei existiert",
                wx.YES_NO | wx.ICON_QUESTION
            )
            if response == wx.NO:
                return
        os.remove(vtt_file_path)

    with open(file_path, "r") as f, open(vtt_file_path, "w") as vtt_file:
        vtt_file.write("WEBVTT\n\n")
        first_line = True
        for line in f:
            if first_line:
                first_line = False # skipping the first line
                continue
            if line.strip().isdigit():  # skipping rows wth digits
                continue
            if "-->" in line:  # converting time format "," into "."
                line = line.replace(",", ".")
            vtt_file.write(line)

    file_name = os.path.basename(vtt_file_path)
    wx.MessageBox(
        f"Die Datei wurde erfolgreich konvertiert und gespeichert als:\n{file_name}",
        "Erfolg",
        wx.OK | wx.ICON_INFORMATION
    )

# Method to handle `Über die App` menu item
def on_about(self, event):
    # Creating a new user-defined window
    frame = wx.Frame(None, title="Über die App", size=(400, 550))
    panel = wx.Panel(frame)
    sizer = wx.BoxSizer(wx.VERTICAL)

    # Adding message
    message = (
        "Name der Anwendung:\n"
        "Dateiablage\n\n"
    )
    message_label = wx.StaticText(panel, label = message)
    sizer.Add(message_label, 0, wx.ALL, 10)

    # Adding logo
    logo = wx.Bitmap("_internal/images/logo.png", wx.BITMAP_TYPE_PNG)
    logo_image = wx.StaticBitmap(panel, bitmap=logo)
    sizer.Add(logo_image, 0, wx.ALL, 10)

    # Adding 2nd message
    message2 = (
        "Versionsnummer:\n"
        "0.2.0 (build 2025-02-04)\n\n"
        "Das Unternehmen:\n"
        "CompuGroup Medical\n\n"
        "**** Beschreibung ****\n"
        "Die Dateiablage ist eine Anwendung zur unkomplizierten Verwaltung von "
        "e-Learning-Inhalten und Aufgaben für das medizinische Personal.\n\n"
        "**** Support ****\n"
        "Bei Fragen oder technischen Problemen kontaktieren Sie bitte unseren Support.\n\n"
    )
    message_label2 = wx.StaticText(panel, label = message2)
    sizer.Add(message_label2, 0, wx.ALL, 10)

    # Adding web link to support page
    support_label = wx.StaticText(panel, label="Weitere Informationen finden Sie auf unserer Support-Seite.")
    support_label.SetForegroundColour((0, 0, 255))  # Link-Farbe (Blau)
    font = support_label.GetFont()
    font.SetUnderlined(True)  # Unterstrichen hinzufügen
    support_label.SetFont(font)

    # Adding click event for the link
    support_label.Bind(wx.EVT_LEFT_DOWN, lambda event: webbrowser.open("https://www.cgm.com/corp_de/unternehmen/kontakt.html"))
    sizer.Add(support_label, 0, wx.ALL, 10)

    # Adding message
    thanks_label = wx.StaticText(panel, label="Vielen Dank, dass Sie unsere Anwendung verwenden!")
    thanks_label.SetForegroundColour((0, 128, 0))  # Grüne Farbe für Freundlichkeit
    sizer.Add(thanks_label, 0, wx.ALL, 10)

    # Setting sizer and display window
    panel.SetSizer(sizer)
    frame.Show()

# Method to handle the Contact menu item
def on_contact(self, event):
    mailto_link = "mailto:support-dateiablage@cgm.com?subject=Supportanfrage&body=Hallo%20Support-Team"
    wx.LaunchDefaultBrowser(mailto_link)

# Method to handle the Check Completeness menu item
def on_check_completeness(self, event):
    check_folder_completeness(self, event)
    
# Method to check the completeness of the folder
def check_folder_completeness(self, event):
    required_extensions = ['.srt', '.vtt', '.wav', '.story', '.mp4', '.trec', '.tscproj']
    present_files = []
    missing_files = []

    # Searching the folder for the required file types
    for ext in required_extensions:
        found = False
        for file in g.file_list:
            if file.endswith(ext):
                found = True
                present_files.append(ext)
                break
        if not found:
            missing_files.append(ext)

    # Creating the message based on the results
    if not missing_files:
        message = "Der Ordner enthält alle benötigten Dateien: " + ", ".join(present_files)
    else:
        message = "Fehlende Dateien:\n" + "\n".join(missing_files)
        message += "\n\nVorhandene Dateien:\n" + "\n".join(present_files)
    wx.MessageBox(message, "Vollständigkeitsprüfung", wx.OK | wx.ICON_INFORMATION)

# Method to handle the Refresh menu item
def on_refresh(self, event):
    # Clearing the ctrl lists
    self.learning_ctrl.DeleteAllItems()
    self.tasks_ctrl.DeleteAllItems()

    # Refreshing the ctrl lists
    try:
        if g.folder_path is not None:
            list_files(self, g.folder_path)
    except Exception as e:
        print(f"Error: {e}")
    try:
        if g.df_elearning is not None:
            display_learning(self, g.df_elearning)
    except Exception as e:
        print(f"Error: {e}")
    try:
        if g.df_tasks is not None:
            on_import_tasks(self, None)
    except Exception as e:
        print(f"Error: {e}")

# Method to handle the Export file list
def on_export(self, event):
    export_docx_multiple_dirs(self, g.file_list)

# Method to export files from multiple directories (non-recursive)
def export_docx_multiple_dirs(self, directories):
    document = Document()
    ignore_list = ["_Protokoll.txt",
                    "_e-Learning_Definition.csv",
                    "_organisatorische_Aufgaben.csv"
]

    # Adding header
    header = document.add_heading('Dateiliste', level=0)
    run = header.runs[0]

    document.add_heading(f"Ordner: {os.path.dirname(directories[0])}", level=1)
    document.add_paragraph("")

    for directory in directories:
        try:
            if os.path.isdir(directory):
                document.add_heading(f"Ordner: {os.path.basename(directory)}", level=2)
            else:
                if not any(item in directory for item in ignore_list):
                    document.add_paragraph(os.path.basename(directory))
        except Exception as e:
            document.add_paragraph(f"Fehler beim Zugriff auf {directory}: {e}")

    # Prompting user for file save location
    save_word_file(document)

# Method to save the Word file
def save_word_file(document):
    with wx.FileDialog(
        None,
        "Speicherort für Exportdatei auswählen",
        defaultFile="Export_Dateiliste.docx",
        wildcard="Word Document (*.docx)|*.docx",
        style=wx.FD_SAVE | wx.FD_OVERWRITE_PROMPT,
    ) as file_dialog:
        if file_dialog.ShowModal() == wx.ID_CANCEL:
            return  # Benutzer hat den Speichervorgang abgebrochen

        save_path = file_dialog.GetPath()
        document.save(save_path)

        wx.MessageBox(
            f"Die Datei wurde erfolgreich exportiert nach:\n{save_path}",
            "Export erfolgreich",
            wx.OK | wx.ICON_INFORMATION,
        )

        # **Datei nach dem Speichern automatisch öffnen**
        try:
            if os.name == "nt":  # Windows
                os.startfile(save_path)
            elif os.name == "posix":  # macOS/Linux
                opener = "open" if sys.platform == "darwin" else "xdg-open"
                subprocess.call([opener, save_path])
        except Exception as e:
            wx.MessageBox(
                f"Die Datei konnte nicht automatisch geöffnet werden:\n{e}",
                "Fehler beim Öffnen",
                wx.OK | wx.ICON_ERROR,
            )

# Method to handle the Exit menu item
def on_exit(self, event):
    self.Close(True)
