import wx # wxPython / Phoenix
import os
import io
import subprocess
import platform
import unicodedata
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MyFrame(wx.Frame):
    def __init__(self, *args, **kw):
        super(MyFrame, self).__init__(*args, **kw)

        # Definition of global variables
        self.file_list = []

        # Creating a menu bar
        menu_bar = wx.MenuBar()

        # Creating the `Datei` menu
        file_menu = wx.Menu()
        import_definition = file_menu.Append(wx.ID_ANY, "&Importiere e-Learning Definition")
        import_tasks = file_menu.Append(wx.ID_ANY, "&Wähle Aufgabenliste")
        browse_item = file_menu.Append(wx.ID_OPEN, "&Wähle Quellverzeichnis")
        
        file_menu.Append(wx.ID_EXIT, "&Beenden")
        menu_bar.Append(file_menu, "&Datei")
        
        # Creating the `Bearbeiten` menu
        edit_menu = wx.Menu()
        export_file_list = edit_menu.Append(wx.ID_ANY, "Exportiere Dateiliste")
        menu_bar.Append(edit_menu, "&Bearbeiten")

        # Setting the menu bar
        self.SetMenuBar(menu_bar)

        # Creating a panel
        panel = wx.Panel(self)

        # Creating the list controls
        self.learning_ctrl = wx.ListCtrl(panel,
                                     style=wx.LC_REPORT
                                     |wx.BORDER_SUNKEN|wx.LIST_ALIGN_SNAP_TO_GRID
                                     )
        self.tasks_ctrl = wx.ListCtrl(panel,
                                     style=wx.LC_ICON
                                     |wx.BORDER_SUNKEN|wx.LIST_ALIGN_SNAP_TO_GRID
                                     )
        self.file_listbox = wx.ListBox(panel)

        # Creating a horizontal box sizer
        hbox = wx.BoxSizer(wx.HORIZONTAL)
        # Adding the list controls to the horizontal box sizer
        hbox.Add(self.learning_ctrl, 1, wx.ALL | wx.EXPAND, 5)
        hbox.Add(self.tasks_ctrl, 1, wx.ALL | wx.EXPAND, 5)

        # Creating a vertical box sizer
        vbox = wx.BoxSizer(wx.VERTICAL)
        # Adding the horizontal box sizer and the file listbox to the vertical box sizer
        vbox.Add(hbox, 1, wx.ALL | wx.EXPAND, 5)
        vbox.Add(self.file_listbox, 1, wx.ALL | wx.EXPAND, 5)
        # Setting the sizer for the frame and fit the panel
        panel.SetSizer(vbox)

        # Binding the Import menu item to the on_import method
        self.Bind(wx.EVT_MENU, self.on_import_csv, import_definition)
        # Binding the Import menu item to the on_import method
        self.Bind(wx.EVT_MENU, self.on_import_excel, import_tasks)
        # Binding the Browse menu item to the on_browse method
        self.Bind(wx.EVT_MENU, self.on_browse, browse_item)
        # Bindung the Export menu item to the on_export method
        self.Bind(wx.EVT_MENU, self.on_export, export_file_list)

        # Binding the list control to the on_item_activated method
        self.learning_ctrl.Bind(wx.EVT_LIST_ITEM_SELECTED, self.on_item_selected)
        # Binding the list control to the on_item_activated method
        self.tasks_ctrl.Bind(wx.EVT_LIST_ITEM_SELECTED, self.on_item_selected)
        # Binding the list control to the on_file_activated method
        self.file_listbox.Bind(wx.EVT_LISTBOX_DCLICK, self.on_file_activated)

    # Method to handle the Export file list
    def on_export(self, event):
        self.export_docx(self.file_list)

    def export_docx(self, data):
        document = Document()

        # Adding header
        document.add_heading('Dateiliste', 0)
        paragraph = document.add_paragraph()
        paragraph.add_run(data)
        document.add_page_break()

        ## Creating a Word file using python-docx as engine
        buffer = io.BytesIO()
        document.save(buffer)
        with open("Export_Dateiliste.docx", "wb") as docx_file:
            docx_file.write(buffer.getvalue())
        buffer.close()

    # Method to handle the Browse menu item
    def on_browse(self, event):
        dialog = wx.DirDialog(None, "Wähle einen Ordner aus:", style=wx.DD_DEFAULT_STYLE | wx.DD_NEW_DIR_BUTTON)
        if dialog.ShowModal() == wx.ID_OK:
            self.folder_path = dialog.GetPath() # folder_path will contain the path of the folder you have selected as string
            self.list_files(self.folder_path)
        dialog.Destroy()

    # Method to handle the Import learning definition
    def on_import_csv(self, event):
        dialog = wx.FileDialog(self, "Importiere e-Learning Definition", wildcard="CSV files (*.csv)|*.csv|All files (*.*)|*.*", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST)
        if dialog.ShowModal() == wx.ID_OK:
            file_path = dialog.GetPath()
            self.import_csv(file_path)
        dialog.Destroy()

    # Method to import the CSV file
    def import_csv(self, file_path):
        try:
            definition_csv = pd.read_csv(file_path)
            self.display_learning(definition_csv)
            wx.MessageBox(f"Datei erfolgreich importiert: {file_path}", "Erfolg", wx.OK | wx.ICON_INFORMATION)
        except Exception as e:
            wx.MessageBox(f"Datei nicht importiert: {e}", "Error", wx.OK | wx.ICON_ERROR)

    # Method to display the data in the learning control
    def display_learning(self, df):
        self.learning_ctrl.ClearAll()
        self.learning_ctrl.AppendColumn("e-Learning Struktur")

        for index, row in df.iterrows():
            text = row.iloc[0]
            level = row.iloc[1]
            indent = ' ' * ((level * 4) - 4)  # Indent based on the level
            self.learning_ctrl.Append([f"{indent}{text}"])

        # Adjusting the column width to fit automatically the content
        self.learning_ctrl.SetColumnWidth(0, wx.LIST_AUTOSIZE)

    # Method to handle the Import tasks excel
    def on_import_excel(self, event):
        dialog = wx.FileDialog(self, "Importiere Aufgabenliste", wildcard="Exceldatei (*.xlsx)|*.xlsx|All files (*.*)|*.*", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST)
        if dialog.ShowModal() == wx.ID_OK:
            file_path = dialog.GetPath()
            self.import_excel(file_path)
        dialog.Destroy()

    # Method to import the Excel file
    def import_excel(self, file_path):
        try:
            data = None
            output = [[]]

            # Reading file as bytes
            with open(file_path, 'rb') as file:
                bytes_data = file.read()

            # Creating pandas dataframe, e.g. sheet_name `1` for sheet no. 2 or "Arzt prozessual"
            data = pd.read_excel(io.BytesIO(bytes_data), sheet_name = 1)

            # Adding from Excel
            for index, row in data.iterrows():
                if len(row) > 1:  # Ensure there are at least 2 columns
                    # Extract relevant columns (adjust the indices as needed)
                    task = str(row.iloc[1]).strip() #zweite Spalte
                    person= str(row.iloc [6]).strip() #siebte Spalte
                    status = str(row.iloc[8]).strip() # neunte Spalte
                    output.append([index, task,person, status])
                else:
                    print(f"Row {index} does not have enough columns: {row}")
            print(output) # Debugging list of lists `output`
            
            # Tasks
            output_df = pd.DataFrame(output, columns = ['ID', 'Aufgabenliste','Verantwortlicher', 'Status'])
            output_df = output_df.set_index('ID')
            output_df = output_df.drop_duplicates(ignore_index = True)
            output_df = output_df.drop(0, axis = 0)
            print(output_df) # Debugging pandas dataframe `output_df`

            self.display_tasks(output_df)
            wx.MessageBox(f"Datei erfolgreich importiert: {file_path}", "Erfolg", wx.OK | wx.ICON_INFORMATION)
        except Exception as e:
            wx.MessageBox(f"Datei nicht importiert: {e}", "Error", wx.OK | wx.ICON_ERROR)

    # Method to display the data in the tasks control
    def display_tasks(self, df):
        self.tasks_ctrl.ClearAll()
        
        #self.tasks_ctrl.SetColumn(1, wx.ListItem('Test')) 
        
        
        for index, row in df.iterrows():
            
          if index==4:  
                self.tasks_ctrl.Append([row.iloc[0]]) 
                self.tasks_ctrl.Append([row.iloc[1]])
                self.tasks_ctrl.Append([row.iloc[2]])
                text="-------------"
                self.tasks_ctrl.Append([text])
                text="-------------"
                self.tasks_ctrl.Append([text])
                self.tasks_ctrl.SetColumnWidth(0, 200)
                self.tasks_ctrl.SetColumnWidth(1, 200)
                self.tasks_ctrl.SetColumnWidth(2, 200)
            
            
                
          
          elif index >= 5:
              
                print(row) # Debugging row data
                   
                #text = row.iloc[0] + " VERANTWORTLICHER "+row.iloc[1]+" STATUS: "+row.iloc[2]
                #self.tasks_ctrl.Append([text])
                
                self.tasks_ctrl.Append([row.iloc[0]]) 
                self.tasks_ctrl.Append([row.iloc[1]])
                self.tasks_ctrl.Append([row.iloc[2]])
                text="-------"
                self.tasks_ctrl.Append([text])
                text="-------"
                self.tasks_ctrl.Append([text])
                
                #self.tasks_ctrl.SetColumnWidth(0, wx.LIST_AUTOSIZE)
                #self.tasks_ctrl.SetColumnWidth(1, 200)
                #self.tasks_ctrl.SetColumnWidth(2, wx.LIST_AUTOSIZE)
                
                
                  
        #self.tasks_ctrl.SetColumnWidth(1,200) #wx.LIST_AUTOSIZE
                
        # Adjusting the column width to fit automatically the content
        #self.tasks_ctrl.SetColumnWidth(0, wx.LIST_AUTOSIZE)

    # Method to handle the list control item activated event
    def on_item_selected(self, event):
        item_index = event.GetIndex()
        item_text = self.learning_ctrl.GetItemText(item_index)
        self.SetTitle(f"Dateiablage - {item_text.strip()}")
        self.list_files(self.folder_path, item_text.strip())
        
    # Method to handle the list control item activated event
    def on_file_activated(self, event):
        file_index = event.GetSelection()
        file_path = self.file_listbox.GetString(file_index)

        if platform.system() == "Windows":
            try:
                os.startfile(f'"{file_path}"')
            except Exception as e:
                wx.MessageBox(f"Datei konnte nicht geöffnet werden: {e}", "Error", wx.OK | wx.ICON_ERROR)
        elif platform.system() == "Darwin":  # macOS
            try:
                subprocess.call(["open", file_path])
            except Exception as e:
                wx.MessageBox(f"Datei konnte nicht geöffnet werden: {e}", "Error", wx.OK | wx.ICON_ERROR)
        else:  # Linux
            try:
                subprocess.call(["xdg-open", file_path])
            except Exception as e:
                wx.MessageBox(f"Datei konnte nicht geöffnet werden: {e}", "Error", wx.OK | wx.ICON_ERROR)
        

    # Method to list the files in the selected folder
    def list_files(self, folder_path, filter_text=None):
        for root, dirs, files in os.walk(folder_path):
            for name in dirs:
                dir_path = os.path.join(root, name)
                if filter_text is None:
                    self.file_list.append(dir_path)
                    files = [
                        f 
                        for f in os.listdir(dir_path)
                    ]
                    self.file_list.extend(os.path.join(dir_path, f) for f in files)
                else:
                    # Adding all subdirectories of the matching directory
                    normalized_filter_text = unicodedata.normalize('NFC', filter_text)
                    normalized_name = unicodedata.normalize('NFC', name)
                    if normalized_filter_text in normalized_name:
                        for sub_root, sub_dirs, sub_files in os.walk(dir_path):
                            for sub_name in sub_dirs:
                                dir_path = os.path.join(sub_root, sub_name)
                                self.file_list.append(os.path.join(sub_root, sub_name))
                                sub_files = [
                                    f 
                                    for f in os.listdir(dir_path)
                                ]
                                self.file_list.extend(os.path.join(dir_path, f) for f in sub_files)

        self.file_listbox.Set(self.file_list)

# Creating the wx App
class MyApp(wx.App):
    def OnInit(self):
        frame = MyFrame(None, title="Dateiablage", size=(1024, 768))
        frame.Show(True)
        return True

# Initializing the wx App
app = MyApp(False)
print("App start")
app.MainLoop()