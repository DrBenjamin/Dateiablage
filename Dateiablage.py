import wx # wxPython / Phoenix
import os
import io
import subprocess
import platform
import unicodedata
import pandas as pd
from docx import Document

class PreferencesPage(wx.PreferencesPage):
    def __init__(self, config):
        super().__init__()
        self.config = config

    def GetName(self):
        return "Einstellungen"

    def CreateWindow(self, parent):
        panel = wx.Panel(parent)
        sizer = wx.BoxSizer(wx.VERTICAL)

        # Adding preference control user settings
        heading_user = wx.StaticText(panel, label="Benutzereinstellungen")
        font = heading_user.GetFont()
        font.PointSize += 2
        heading_user.SetFont(font)
        sizer.Add(heading_user, 0, wx.ALL, 5)
        # User choice
        self.user_choice = wx.Choice(panel, choices=["Asher", "Benjamin", "Larisa", "Listan", "Mahsa", "Marko", "Sandra"])
        sizer.Add(self.user_choice, 0, wx.ALL, 5)
        # Load saved state
        user_state = self.config.Read("user_choice", "Asher")
        self.user_choice.SetStringSelection(user_state)
        # Bind event to save state
        self.user_choice.Bind(wx.EVT_CHOICE, self.on_user_choice)

        # Adding preference control csv import
        heading_csv = wx.StaticText(panel, label="CSV Import Einstellungen")
        font = heading_csv.GetFont()
        font.PointSize += 2
        heading_csv.SetFont(font)
        sizer.Add(heading_csv, 0, wx.ALL, 5)
        # CSV Import checkbox
        self.csv_checkbox = wx.CheckBox(panel, label="CSV Import")
        sizer.Add(self.csv_checkbox, 0, wx.ALL, 5)
        # Load saved state
        csv_state = self.config.ReadBool("csv_import_enabled", False)
        self.csv_checkbox.SetValue(csv_state)
        # Bind event to save state
        self.csv_checkbox.Bind(wx.EVT_CHECKBOX, self.on_csv_checkbox)

        # Adding preference control srt converter
        heading = wx.StaticText(panel, label="SRT Konverter Einstellungen")
        font = heading.GetFont()
        font.PointSize += 2 
        heading.SetFont(font)
        sizer.Add(heading, 0, wx.ALL, 5)
        # SRT Konverter checkbox
        self.srt_checkbox = wx.CheckBox(panel, label="SRT Konverter")
        sizer.Add(self.srt_checkbox, 0, wx.ALL, 5)
        # Load saved state
        srt_state = self.config.ReadBool("srt_converter_enabled", False)
        self.srt_checkbox.SetValue(srt_state)
        # Bind event to save state
        self.srt_checkbox.Bind(wx.EVT_CHECKBOX, self.on_srt_checkbox)

        panel.SetSizer(sizer)
        return panel

    # Method to handle the Preferences page user choice
    def on_user_choice(self, event):
        self.config.Write("user_choice", self.user_choice.GetString(self.user_choice.GetSelection()))
        self.config.Flush()

    # Method to handle the Preferences page csv checkbox
    def on_csv_checkbox(self, event):
        self.config.WriteBool("csv_import_enabled", self.csv_checkbox.IsChecked())
        self.config.Flush()

    # Method to handle the Preferences page srt checkbox
    def on_srt_checkbox(self, event):
        self.config.WriteBool("srt_converter_enabled", self.srt_checkbox.IsChecked())
        self.config.Flush()

class MyFrame(wx.Frame):
    def __init__(self, *args, **kw):
        super(MyFrame, self).__init__(*args, **kw)

        # Initialize config
        self.config = wx.Config("Dateiablage")

        # Definition of global variables
        self.file_list = []
        self.definition_csv = None
        self.file_path_tasks = None
        self.folder_path = None

        # Creating a menu bar
        menu_bar = wx.MenuBar()

        # Creating the `Datei` menu
        file_menu = wx.Menu()
        import_definition = file_menu.Append(wx.ID_ANY, "&Importiere e-Learning Definition")
        import_tasks = file_menu.Append(wx.ID_ANY, "&Wähle Aufgabenliste")
        browse_item = file_menu.Append(wx.ID_OPEN, "&Wähle Quellverzeichnis")
        exit_app = file_menu.Append(wx.ID_EXIT, "&Beenden")
        menu_bar.Append(file_menu, "&Datei")
        
        # Creating the `Bearbeiten` menu
        edit_menu = wx.Menu()
        export_file_list = edit_menu.Append(wx.ID_ANY, "Exportiere Dateiliste")
        refresh_ctrl_lists = edit_menu.Append(wx.ID_ANY, "Aktualisiere")
        preferences = edit_menu.Append(wx.ID_ANY, "Einstellungen")
        menu_bar.Append(edit_menu, "&Bearbeiten")

        # Setting the menu bar
        self.SetMenuBar(menu_bar)

        # Creating a panel
        panel = wx.Panel(self)

        # Creating the list controls
        self.learning_ctrl = wx.ListCtrl(panel,
                                     style=wx.LC_REPORT
                                     |wx.BORDER_SUNKEN|wx.LIST_ALIGN_SNAP_TO_GRID
                                     )
        self.tasks_ctrl = wx.ListCtrl(panel,
                                     style=wx.LC_ICON
                                     |wx.BORDER_SUNKEN|wx.LIST_ALIGN_SNAP_TO_GRID
                                     )
        self.file_listbox = wx.ListBox(panel)

        # Creating a horizontal box sizer
        hbox = wx.BoxSizer(wx.HORIZONTAL)
        # Adding the list controls to the horizontal box sizer
        hbox.Add(self.learning_ctrl, 1, wx.ALL | wx.EXPAND, 5)
        hbox.Add(self.tasks_ctrl, 1, wx.ALL | wx.EXPAND, 5)

        # Creating a vertical box sizer
        vbox = wx.BoxSizer(wx.VERTICAL)
        # Adding the horizontal box sizer and the file listbox to the vertical box sizer
        vbox.Add(hbox, 1, wx.ALL | wx.EXPAND, 5)
        vbox.Add(self.file_listbox, 1, wx.ALL | wx.EXPAND, 5)
        # Setting the sizer for the frame and fit the panel
        panel.SetSizer(vbox)

        # Binding the Import menu item to the on_import method
        self.Bind(wx.EVT_MENU, self.on_import_csv, import_definition)
        # Binding the Import menu item to the on_import method
        self.Bind(wx.EVT_MENU, self.on_import_excel, import_tasks)
        # Binding the Browse menu item to the on_browse method
        self.Bind(wx.EVT_MENU, self.on_browse, browse_item)
        # Bindung the Export menu item to the on_export method
        self.Bind(wx.EVT_MENU, self.on_export, export_file_list)
        # Binding the Exit menu item to the on_exit method
        self.Bind(wx.EVT_MENU, self.on_exit, exit_app)
        # Binding the Refresh menu item to the on_refresh method
        self.Bind(wx.EVT_MENU, self.on_refresh, refresh_ctrl_lists)
        # Binding the Preferences menu item to the on_preferences method
        self.Bind(wx.EVT_MENU, self.on_preferences, preferences)

        # Binding the list control to the on_item_activated method
        self.learning_ctrl.Bind(wx.EVT_LIST_ITEM_SELECTED, self.on_item_selected)
        # Binding the list control to the on_item_activated method
        self.tasks_ctrl.Bind(wx.EVT_LIST_ITEM_SELECTED, self.on_item_selected)
        # Binding the list control to the on_file_activated method
        self.file_listbox.Bind(wx.EVT_LISTBOX_DCLICK, self.on_file_activated)
    # Method to handle the Refresh menu item

    def on_refresh(self, event):
        # Clear the ctrl lists
        self.learning_ctrl.DeleteAllItems()
        self.tasks_ctrl.DeleteAllItems()
        
        # Refresh the ctrl lists
        try:
            if self.folder_path is not None:
                self.list_files(self.folder_path)
        except Exception as e:
            print(f"Error: {e}")
        try:
            if self.definition_csv is not None:
                self.display_learning(self.definition_csv)
        except Exception as e:
            print(f"Error: {e}")
        try:
            if self.file_path_tasks is not None:
                self.import_excel(self.file_path_tasks)
        except Exception as e:
            print(f"Error: {e}")

    # Method to handle the Preferences menu item
    def on_preferences(self, event):
        dialog = wx.PreferencesEditor()
        dialog.AddPage(PreferencesPage(self.config))
        dialog.Show(self)

    # Method to handle the Export file list
    def on_export(self, event):
        self.export_docx(self.file_list)

    # Method to handle the Exit menu item
    def on_exit(self, event):
        self.Close(True)

    # Method to export the file list to a Word document
    def export_docx(self, data):
        document = Document()

        # Adding header
        document.add_heading('Dateiliste', 0)
        paragraph = document.add_paragraph()
        paragraph.add_run(data)
        document.add_page_break()

        ## Creating a Word file using python-docx as engine
        buffer = io.BytesIO()
        document.save(buffer)
        with open("Export_Dateiliste.docx", "wb") as docx_file:
            docx_file.write(buffer.getvalue())
        buffer.close()

    # Method to handle the Import learning definition
    def on_import_csv(self, event):
        if self.config.ReadBool("csv_import_enabled", True):
            dialog = wx.FileDialog(self, "Importiere e-Learning Definition", wildcard="CSV files (*.csv)|*.csv|All files (*.*)|*.*", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST)
            if dialog.ShowModal() == wx.ID_OK:
                file_path = dialog.GetPath()
                self.import_csv(file_path)
            dialog.Destroy()
        else:
            wx.MessageBox("CSV Import ist deaktiviert (siehe Einstellungen)", "Information", wx.OK | wx.ICON_WARNING)

    # Method to import the CSV file
    def import_csv(self, file_path):
        try:
            self.definition_csv = pd.read_csv(file_path)
            self.display_learning(self.definition_csv)
            wx.MessageBox(f"Datei erfolgreich importiert: {file_path}", "Erfolg", wx.OK | wx.ICON_INFORMATION)
        except Exception as e:
            wx.MessageBox(f"Datei nicht importiert: {e}", "Error", wx.OK | wx.ICON_ERROR)

    # Method to display the data in the learning control
    def display_learning(self, df):
        self.learning_ctrl.ClearAll()
        self.learning_ctrl.AppendColumn("e-Learning Struktur")

        for index, row in df.iterrows():
            text = row.iloc[0]
            level = row.iloc[1]
            indent = ' ' * ((level * 4) - 4)  # Indent based on the level
            self.learning_ctrl.Append([f"{indent}{text}"])

        # Adjusting the column width to fit automatically the content
        self.learning_ctrl.SetColumnWidth(0, wx.LIST_AUTOSIZE)

    # Method to handle the Import tasks excel
    def on_import_excel(self, event):
        dialog = wx.FileDialog(self, "Importiere Aufgabenliste", wildcard="Exceldatei (*.xlsx)|*.xlsx|All files (*.*)|*.*", style=wx.FD_OPEN | wx.FD_FILE_MUST_EXIST)
        if dialog.ShowModal() == wx.ID_OK:
            self.file_path_tasks = dialog.GetPath()
            self.import_excel(self.file_path_tasks)
        dialog.Destroy()

    # Method to import the Excel file
    def import_excel(self, file_path):
        try:
            data = None
            output = [[]]

            # Reading file as bytes
            with open(file_path, 'rb') as file:
                bytes_data = file.read()

            # Creating pandas dataframe, e.g. sheet_name `1` for sheet no. 2 or "Arzt prozessual"
            data = pd.read_excel(io.BytesIO(bytes_data), sheet_name = 1)
            print("Excel: ", data)
            
            # Tasks from Excel
            for index, row in data.iterrows():
                if len(row) > 1:  # Ensure there are at least 2 columns
                    # Extract relevant columns (adjust the indices as needed)
                    task = str(row.iloc[1]).strip()
                    user = str(row.iloc[6]).strip()
                    status = str(row.iloc[8]).strip()
                    if index >= 5:
                        if self.config.Read("user_choice") == user:
                            output.append([index, task, user, status])
                else:
                    print(f"Row {index} does not have enough columns: {row}")
            print("Output: ", output) # Debugging list of lists `output`
            if output == [[]]:
                output_df = pd.DataFrame({
                                                'ID': "0",
                                                'Aufgabenliste': ['leer'],
                                                'Verantwortlicher': ['leer'],
                                                'Status': ['leer']
                                    })
                output_df = output_df.set_index('ID')
            else:
                output_df = pd.DataFrame(output, columns = ['ID', 'Aufgabe', 'Verantwortlicher', 'Status'])
                output_df = output_df.set_index('ID')
                output_df = output_df.drop_duplicates(ignore_index = True)
                output_df = output_df.drop(0, axis = 0)
            print("Pandas df: ", len(output_df)) # Debugging pandas dataframe `output_df`

            self.display_tasks(output_df)
            wx.MessageBox(f"Datei erfolgreich importiert: {file_path}", "Erfolg", wx.OK | wx.ICON_INFORMATION)
        except Exception as e:
            wx.MessageBox(f"Datei nicht importiert: {e}", "Error", wx.OK | wx.ICON_ERROR)

    # Method to display the data in the tasks control
    def display_tasks(self, df):
        self.tasks_ctrl.ClearAll()
        self.tasks_ctrl.AppendColumn("Aufgabenliste")
        for index, row in df.iterrows():
            print(row) # Debugging row data
            text = "Aufgabe: " + row.iloc[0] + " Verantwortlicher: " + row.iloc[1] + " Status: " + row.iloc[2]
            self.tasks_ctrl.Append([text])

        # Adjusting the column width to fit automatically the content
        self.tasks_ctrl.SetColumnWidth(0, 200)
        #self.tasks_ctrl.SetColumnWidth(0, wx.LIST_AUTOSIZE)

    # Method to handle the Browse menu item
    def on_browse(self, event):
        dialog = wx.DirDialog(None, "Wähle einen Ordner aus:", style=wx.DD_DEFAULT_STYLE | wx.DD_NEW_DIR_BUTTON)
        if dialog.ShowModal() == wx.ID_OK:
            self.folder_path = dialog.GetPath() # folder_path will contain the path of the folder you have selected as string
            self.list_files(self.folder_path)
        dialog.Destroy()
        
    # Method to handle the list control item activated event
    def on_item_selected(self, event):
        item_index = event.GetIndex()
        item_text = self.learning_ctrl.GetItemText(item_index)
        self.SetTitle(f"Dateiablage - {item_text.strip()}")
        self.list_files(self.folder_path, item_text.strip())
        
    # Method to handle the list control item activated event
    def on_file_activated(self, event):
        file_index = event.GetSelection()
        file_path = self.file_listbox.GetString(file_index)

        if platform.system() == "Windows":
            try:
                os.startfile(f'"{file_path}"')
            except Exception as e:
                wx.MessageBox(f"Datei konnte nicht geöffnet werden: {e}", "Error", wx.OK | wx.ICON_ERROR)
        elif platform.system() == "Darwin":  # macOS
            try:
                subprocess.call(["open", file_path])
            except Exception as e:
                wx.MessageBox(f"Datei konnte nicht geöffnet werden: {e}", "Error", wx.OK | wx.ICON_ERROR)
        else:  # Linux
            try:
                subprocess.call(["xdg-open", file_path])
            except Exception as e:
                wx.MessageBox(f"Datei konnte nicht geöffnet werden: {e}", "Error", wx.OK | wx.ICON_ERROR)

    # Method to list the files in the selected folder
    def list_files(self, folder_path, filter_text=None):
        # Clear the existing file list
        self.file_list = []
        for root, dirs, files in os.walk(folder_path):
            for name in dirs:
                dir_path = os.path.join(root, name)
                if filter_text is None:
                    self.file_list.append(dir_path)
                    files = [
                        f 
                        for f in os.listdir(dir_path)
                    ]
                    self.file_list.extend(os.path.join(dir_path, f) for f in files)
                else:
                    # Adding all subdirectories of the matching directory
                    normalized_filter_text = unicodedata.normalize('NFC', filter_text)
                    normalized_name = unicodedata.normalize('NFC', name)
                    if normalized_filter_text in normalized_name:
                        for sub_root, sub_dirs, sub_files in os.walk(dir_path):
                            for sub_name in sub_dirs:
                                dir_path = os.path.join(sub_root, sub_name)
                                self.file_list.append(os.path.join(sub_root, sub_name))
                                sub_files = [
                                    f 
                                    for f in os.listdir(dir_path)
                                ]
                                self.file_list.extend(os.path.join(dir_path, f) for f in sub_files)
        self.file_listbox.Set(self.file_list)

# Creating the wx App
class MyApp(wx.App):
    def OnInit(self):
        frame = MyFrame(None, title="Dateiablage", size=(1024, 768))
        frame.Show(True)
        return True

# Initializing the wx App
app = MyApp(False)
print("App start")
app.MainLoop()